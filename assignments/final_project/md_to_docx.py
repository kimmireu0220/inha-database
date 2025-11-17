#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def parse_markdown_to_docx(md_file_path, docx_file_path):
    """마크다운 파일을 워드 문서로 변환"""
    
    # 워드 문서 생성
    doc = Document()
    
    # 기본 스타일 설정 (한글 폰트)
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    # 마크다운 파일 읽기
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 빈 줄 처리
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        
        # 제목 처리
        if line.startswith('# '):
            # H1
            p = doc.add_heading(line[2:].strip(), level=1)
            i += 1
        elif line.startswith('## '):
            # H2
            p = doc.add_heading(line[3:].strip(), level=2)
            i += 1
        elif line.startswith('### '):
            # H3
            p = doc.add_heading(line[4:].strip(), level=3)
            i += 1
        elif line.startswith('#### '):
            # H4
            p = doc.add_heading(line[5:].strip(), level=4)
            i += 1
        elif line.startswith('##### '):
            # H5
            p = doc.add_heading(line[6:].strip(), level=5)
            i += 1
        elif line.startswith('###### '):
            # H6
            p = doc.add_heading(line[7:].strip(), level=6)
            i += 1
        # 구분선 처리
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)
            i += 1
        # 코드 블록 처리
        elif line.startswith('```'):
            code_language = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # 닫는 ``` 건너뛰기
            
            # 코드 블록 추가
            code_text = ''.join(code_lines).rstrip()
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            # 코드 스타일 적용 (고정폭 폰트)
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
        # 리스트 처리
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # 강조 처리
            p = doc.add_paragraph(text, style='List Bullet')
            process_inline_formatting(p, text)
            i += 1
        elif re.match(r'^\d+\.\s', line):
            # 번호 리스트
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            process_inline_formatting(p, text)
            i += 1
        # 일반 텍스트 처리
        else:
            # 여러 줄을 하나의 문단으로 합치기
            paragraph_text = []
            while i < len(lines):
                current_line = lines[i].rstrip()
                # 제목, 코드 블록, 리스트 시작이면 중단
                if (current_line.startswith('#') or 
                    current_line.startswith('```') or
                    current_line.startswith('- ') or
                    current_line.startswith('* ') or
                    re.match(r'^\d+\.\s', current_line) or
                    current_line.startswith('---')):
                    break
                if current_line.strip():
                    paragraph_text.append(current_line)
                elif paragraph_text:  # 빈 줄이지만 이미 텍스트가 있으면 중단
                    break
                i += 1
            
            if paragraph_text:
                text = ' '.join(paragraph_text)
                p = doc.add_paragraph()
                process_inline_formatting(p, text)
            else:
                i += 1
    
    # 문서 저장
    doc.save(docx_file_path)
    print(f"워드 문서가 생성되었습니다: {docx_file_path}")

def process_inline_formatting(paragraph, text):
    """인라인 포맷팅 처리 (강조, 코드, 링크, 이미지)"""
    
    # 이미지 처리
    img_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    img_matches = list(re.finditer(img_pattern, text))
    
    # 코드 인라인 처리
    code_pattern = r'`([^`]+)`'
    code_matches = list(re.finditer(code_pattern, text))
    
    # 강조 처리 (**bold**, *italic*)
    bold_pattern = r'\*\*([^*]+)\*\*'
    bold_matches = list(re.finditer(bold_pattern, text))
    
    italic_pattern = r'\*([^*]+)\*'
    italic_matches = list(re.finditer(italic_pattern, text))
    
    # 모든 매치를 위치순으로 정렬
    all_matches = []
    for m in img_matches:
        all_matches.append(('img', m.start(), m.end(), m))
    for m in code_matches:
        all_matches.append(('code', m.start(), m.end(), m))
    for m in bold_matches:
        all_matches.append(('bold', m.start(), m.end(), m))
    for m in italic_matches:
        all_matches.append(('italic', m.start(), m.end(), m))
    
    all_matches.sort(key=lambda x: x[1])
    
    # 텍스트를 순차적으로 처리
    last_pos = 0
    for match_type, start, end, match in all_matches:
        # 이전 텍스트 추가
        if start > last_pos:
            normal_text = text[last_pos:start]
            if normal_text:
                run = paragraph.add_run(normal_text)
        
        # 매치된 부분 처리
        if match_type == 'img':
            # 이미지는 텍스트로 표시 (경로 포함)
            img_text = f"[이미지: {match.group(1)} - {match.group(2)}]"
            run = paragraph.add_run(img_text)
            run.italic = True
            run.font.color.rgb = RGBColor(0, 0, 255)  # 파란색
        elif match_type == 'code':
            # 인라인 코드
            code_text = match.group(1)
            run = paragraph.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 0, 128)  # 보라색
        elif match_type == 'bold':
            # 굵게
            bold_text = match.group(1)
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif match_type == 'italic':
            # 기울임
            italic_text = match.group(1)
            run = paragraph.add_run(italic_text)
            run.italic = True
        
        last_pos = end
    
    # 남은 텍스트 추가
    if last_pos < len(text):
        normal_text = text[last_pos:]
        if normal_text:
            run = paragraph.add_run(normal_text)

if __name__ == '__main__':
    md_file = '최종프로젝트_보고서.md'
    docx_file = '최종프로젝트_보고서.docx'
    
    if os.path.exists(md_file):
        parse_markdown_to_docx(md_file, docx_file)
    else:
        print(f"파일을 찾을 수 없습니다: {md_file}")

